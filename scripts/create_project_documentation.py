from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Altheria_FTS_Project_Documentation.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 33, 45)
MUTED = RGBColor(90, 101, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Altheria FTS v3 Project Documentation")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Forensic Tactical Simulation System - Architecture, Implementation, and Workflow Guide")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    rows = [
        ("Project", "Altheria Forensic Tactical Simulation System v3"),
        ("Document date", date.today().strftime("%d %B %Y")),
        ("Codebase", str(ROOT)),
        ("Primary purpose", "Case-based forensic assessment, evidence recording, adaptive questioning, and report generation."),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    add_callout(
        doc,
        "Important scope note",
        "This document describes the current local project implementation: React/Vite frontend, Express backend, browser speech and media APIs, local JSON persistence, recordings, and PDF reporting. Browser-level camera and microphone access still depends on the user's OS and browser permissions.",
    )


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_label_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[idx])
    set_table_geometry(table, widths)
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(40, 45, 52)


def build_document():
    doc = Document()
    configure_styles(doc)

    header = doc.sections[0].header.paragraphs[0]
    header.text = "Altheria FTS v3 - Project Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = doc.sections[0].footer.paragraphs[0]
    footer.add_run("Page ")
    add_page_number(footer)

    add_title(doc)

    doc.add_heading("1. Executive Overview", level=1)
    doc.add_paragraph(
        "Altheria FTS v3 is a full-stack forensic assessment and tactical simulation application. It lets an investigator verify a case, start an assessment session, ask adaptive live questions, capture the subject's spoken answers without a manual send button, record evidence, and generate case-specific forensic reports."
    )
    add_bullets(doc, [
        "The application is designed for local development and demonstration with a browser client and Node.js backend.",
        "The assessment flow supports English and Hindi speech interaction through browser text-to-speech and speech recognition APIs.",
        "Reports and recordings are filtered by Case ID so an investigator sees only evidence linked to the selected case.",
        "Recordings and reports are stored locally under the server uploads folder, while case/session metadata is persisted in a JSON data store.",
    ])

    doc.add_heading("2. Technology Stack", level=1)
    add_matrix(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ("Frontend", "React 18, JavaScript/JSX, Vite 5", "Single-page application, routing, assessment UI, reports UI, recordings UI."),
            ("Styling", "Tailwind CSS, custom CSS", "Responsive interface, panels, drawers, assessment layout, and command hub screens."),
            ("Browser APIs", "SpeechSynthesis, SpeechRecognition/webkitSpeechRecognition, MediaDevices, MediaRecorder, Web Audio, optional FaceDetector", "Hands-free questioning, voice capture, camera capture, evidence recording, and browser-side signal extraction."),
            ("Backend", "Node.js, Express.js, ES Modules", "REST API, authentication, case/session lifecycle, recordings, reports, and audit routes."),
            ("Security", "JWT, bcryptjs, express-rate-limit, CORS", "Investigator login, protected APIs, password hashing, rate limiting, and local client/server access."),
            ("Reporting", "PDFKit", "Generated forensic PDF reports with case, transcript, voice, camera, and recording sections."),
            ("Storage", "Local JSON plus upload folders", "Persistence for cases, sessions, reports, recordings, and audit log entries."),
            ("Optional AI", "Local Ollama API with rule fallback", "Adaptive forensic follow-up questions when a local model is available."),
        ],
        [1600, 3100, 4660],
    )

    doc.add_heading("3. Repository Structure", level=1)
    add_code_block(doc, [
        "altheria-v3-fixed/",
        "  package.json                     Root scripts for install, dev, and build",
        "  client/                          React/Vite frontend",
        "    src/App.jsx                    App routing",
        "    src/components/                Gateway, command hub, assessment, reports, intake",
        "    src/context/AppContext.jsx     Auth/session context",
        "    src/services/api.js            Central frontend API client",
        "  server/                          Express backend",
        "    src/index.js                   API server, middleware, WebSocket binding",
        "    src/db.js                      Local JSON-backed persistence",
        "    src/routes/                    Auth, cases, sessions, interrogation, recordings, reports, audit",
        "    src/services/aiInterrogator.js Adaptive questions and response analysis",
        "    src/services/pdfReportService.js PDF report generation",
        "  uploads/                         Generated data, recordings, and reports",
    ])

    doc.add_heading("4. Main User Workflow", level=1)
    add_numbered(doc, [
        "Investigator logs in through the Gateway screen using a badge ID and passcode.",
        "The Command Hub opens the operational workspace with intake, reports, and recordings drawers.",
        "A case is verified or registered, then a new assessment session is created for that Case ID.",
        "The Assessment Enclave loads the session, asks the opening question aloud, and starts listening.",
        "The subject answers by voice. After five seconds of silence, the browser submits the captured text automatically.",
        "The backend records the answer, analyzes linguistic stress, generates the next question, and returns it.",
        "The frontend speaks the next question aloud and repeats the same listen-and-submit cycle.",
        "At session completion, recordings are uploaded and a PDF report is generated for the case.",
        "Reports and recordings can later be opened by entering the Case ID in the Reports/History or Recordings panels.",
    ])

    doc.add_heading("5. Frontend Implementation", level=1)
    add_label_table(doc, [
        ("Framework", "React 18 with Vite 5."),
        ("Language", "JavaScript and JSX."),
        ("Routing", "react-router-dom controls login, command hub, case flows, and assessment screens."),
        ("API layer", "client/src/services/api.js centralizes fetch calls and JWT headers."),
        ("State management", "React hooks and AppContext maintain investigator, token, and navigation state."),
        ("Primary assessment component", "client/src/components/AssessmentEnclave.jsx handles speech, listening, camera, recording, live metrics, transcript, report generation, and completion actions."),
    ])

    doc.add_heading("6. Backend Implementation", level=1)
    add_label_table(doc, [
        ("Runtime", "Node.js with Express.js using ES Modules."),
        ("Entry point", "server/src/index.js creates the REST server, middleware, routes, health check, and WebSocket server."),
        ("Authentication", "server/src/routes/auth.js validates badge credentials and issues JWT tokens."),
        ("Persistence", "server/src/db.js stores data in maps and writes them to uploads/data/store.json."),
        ("Recording upload", "server/src/routes/recordings.js accepts multipart uploads through Busboy and stores files under uploads/recordings."),
        ("PDF reports", "server/src/services/pdfReportService.js builds case-specific forensic PDF reports with PDFKit."),
        ("Interrogation engine", "server/src/services/aiInterrogator.js generates opening/follow-up questions and analyzes subject responses."),
    ])

    doc.add_heading("7. Assessment Section", level=1)
    doc.add_paragraph(
        "The assessment section is the heart of the project. Its purpose is to run a complete hands-free questioning session: speak the question, listen to the subject, detect silence, submit the answer, request the next question, and continue until the backend ends the assessment."
    )
    add_matrix(
        doc,
        ["Feature", "Code/API Used", "What It Does"],
        [
            ("Question spoken aloud", "window.speechSynthesis, SpeechSynthesisUtterance", "Reads each returned question aloud in the selected language."),
            ("English/Hindi language", "Language mode en/hi, BCP-47 en-IN/hi-IN", "Selects speech synthesis and recognition language for the assessment."),
            ("Live answer capture", "SpeechRecognition or webkitSpeechRecognition", "Continuously listens for spoken answers without a mic/send button."),
            ("Five-second silence submit", "Silence timer in AssessmentEnclave.jsx", "If no speech arrives for five seconds, the current answer is treated as complete."),
            ("Next question loop", "interrogationAPI.respond(sessionId, text, language, context)", "Sends the answer and analysis context to the backend, then receives the next question."),
            ("Completion", "Report generation and navigation actions", "Uploads evidence, generates a report, and offers report/command hub actions."),
        ],
        [2100, 3000, 4260],
    )

    doc.add_heading("8. Live Questioning and Answering Logic", level=1)
    add_numbered(doc, [
        "The frontend opens the session with interrogationAPI.open(sessionId, language).",
        "The returned question is stored as the current prompt and passed to speech synthesis.",
        "When speech synthesis ends, the recognizer starts listening automatically.",
        "Interim and final speech recognition results are accumulated into the current answer buffer.",
        "Every recognized phrase resets the five-second silence timer.",
        "When the timer expires, the latest heard answer is submitted to the backend.",
        "The backend appends the answer to the transcript and returns a new adaptive question.",
        "The frontend speaks that next question and starts listening again.",
    ])
    add_callout(
        doc,
        "Why no send button is needed",
        "The answer-completion trigger is silence, not a click. This matches the required flow: the subject speaks naturally, pauses for five seconds, and the system moves forward by itself.",
    )

    doc.add_heading("9. English and Hindi Support", level=1)
    doc.add_paragraph(
        "The frontend and backend both carry language information. English uses en-IN and Hindi uses hi-IN where the browser supports it. The backend language profile controls localized opening prompts, adaptive probe templates, and Hindi stress/sentiment keyword matching."
    )
    add_bullets(doc, [
        "Frontend speech output: browser SpeechSynthesis chooses the closest available English or Hindi voice.",
        "Frontend speech input: browser SpeechRecognition receives en-IN or hi-IN as the recognizer language.",
        "Backend question generation: aiInterrogator normalizes language and uses matching question profiles.",
        "Backend analysis: stress indicators and sentiment words include English and Hindi patterns.",
        "Fallback behavior: if a browser lacks Hindi speech recognition support, the application cannot force support; Chrome/Edge are recommended.",
    ])

    doc.add_heading("10. Recording and Evidence Capture", level=1)
    doc.add_paragraph(
        "The recording design uses browser media streams and uploads finished blobs to the backend. The system keeps one assessment session connected to its recording metadata so reports and the recordings panel can show evidence by Case ID."
    )
    add_matrix(
        doc,
        ["Recording Type", "Browser API", "Purpose"],
        [
            ("Voice recording", "MediaDevices.getUserMedia({ audio: true }) + MediaRecorder", "Captures the subject's voice answer stream for the session."),
            ("Master recording", "Audio tracks plus available video tracks + MediaRecorder", "Captures combined session evidence when the camera is available."),
            ("Recording metadata", "FormData upload fields", "Stores modality, label, transcript, voice metrics, camera metrics, and analysis summary."),
            ("Recording archive", "recordingsAPI.list(caseId)", "Shows only recordings linked to the searched Case ID."),
        ],
        [2200, 3300, 3860],
    )
    add_callout(
        doc,
        "Browser TTS limitation",
        "The spoken question comes from browser text-to-speech. Browsers usually do not route their own speaker output into the microphone recording. To preserve question evidence, the project stores the full transcript and question text in metadata and reports.",
    )

    doc.add_heading("11. Camera and Behavior Analysis", level=1)
    doc.add_paragraph(
        "The camera pipeline requests video permission through getUserMedia and displays a live preview during assessment. The project also extracts browser-side behavior signals such as motion, brightness/balance changes, and optional face detection when the browser exposes FaceDetector."
    )
    add_bullets(doc, [
        "Camera preview depends on browser and operating system permissions. The app can request access and retry, but it cannot override blocked OS/browser privacy settings.",
        "The Retry Camera action attempts to reacquire the video stream without breaking the active voice assessment flow.",
        "Behavior metrics are included in each submitted answer as behaviorMetrics and later appear in reports.",
        "The analysis is useful as a project-side behavioral indicator, not as medical-grade biometric eye tracking.",
    ])

    doc.add_heading("12. Interrogation Engine and Analysis", level=1)
    doc.add_paragraph(
        "The backend interrogation service controls adaptive questions and textual/metric analysis. It can call a local Ollama model when available; otherwise it uses deterministic localized prompts so the project still works without paid APIs or internet access."
    )
    add_label_table(doc, [
        ("Opening question", "getOpeningStatement creates the first formal psychological assessment prompt."),
        ("Follow-up question", "getNextQuestion creates a follow-up based on transcript, case charge, subject name, and last response."),
        ("Maximum responses", "MAX_INTERROGATION_RESPONSES controls the session length; default is 8 subject responses."),
        ("Stress patterns", "Regex patterns detect evasion, deflection, legal invocation, distancing, minimization, and other linguistic markers."),
        ("Sentiment analysis", "Keyword lexicons score positive, negative, or neutral sentiment."),
        ("Forensic score", "Combines linguistic severity, voice stress, behavior stress, and sentiment weight into a 0-100 score."),
    ])

    doc.add_heading("13. API Endpoints", level=1)
    add_matrix(
        doc,
        ["Area", "Endpoint", "Purpose"],
        [
            ("Auth", "POST /api/auth/login", "Validate investigator badge and passcode, return JWT."),
            ("Cases", "GET /api/cases/:caseId", "Verify a case before opening reports, recordings, or assessment."),
            ("Cases", "POST /api/cases", "Register a new case from intake."),
            ("Sessions", "POST /api/sessions", "Start an assessment session for a Case ID."),
            ("Sessions", "GET /api/sessions/:sessionId", "Load an existing session."),
            ("Interrogation", "POST /api/interrogation/:sessionId/open", "Create the first question."),
            ("Interrogation", "POST /api/interrogation/:sessionId/respond", "Submit answer, analysis context, and receive next question."),
            ("Recordings", "POST /api/recordings", "Upload voice/master evidence files and metadata."),
            ("Recordings", "GET /api/recordings?caseId=...", "List recordings for one case."),
            ("Reports", "POST /api/reports/generate", "Generate a PDF report from a completed session."),
            ("Reports", "GET /api/reports?caseId=...", "List reports for one case."),
            ("Audit", "GET /api/audit", "Read recent audit trail entries."),
            ("Health", "GET /api/health", "Confirm backend service status."),
        ],
        [1500, 3300, 4560],
    )

    doc.add_heading("14. Reports Section", level=1)
    doc.add_paragraph(
        "Reports are generated with PDFKit on the backend. The report generator builds a structured forensic document from the session, transcript, analysis results, recording metadata, and case information."
    )
    add_bullets(doc, [
        "Case Information: Case ID, subject, charge, status, investigator, and timestamps.",
        "Executive Summary: overall findings and session-level summary.",
        "Stress Flags: detected indicators and answer-level severity.",
        "Answer-Level Forensic Analysis: each answer with score, sentiment, voice state, camera stress, and indicators.",
        "Voice Analysis: sentiment and prosody/stress details.",
        "Camera Analysis: behavior and face movement metrics when camera data is available.",
        "Recording Evidence Archive: linked recording files and metadata.",
        "Full Transcript: questions and answers from the session.",
    ])

    doc.add_heading("15. Reports and Recordings Filtering", level=1)
    doc.add_paragraph(
        "The ReportsPanel now asks for a Case ID before showing history. This prevents reports and recordings from different people from appearing together. Reports and recordings are also separated by panel mode, so opening Reports shows reports only, and opening Recordings shows recording evidence only."
    )
    add_label_table(doc, [
        ("Report mode", "User enters Case ID, app verifies the case, then lists only reports for that case."),
        ("Recording mode", "User enters Case ID, app verifies the case, then lists only recordings for that case."),
        ("Download behavior", "Report download opens the PDF report; recording playback/download stays in the recordings panel."),
        ("Navigation", "CommandHub reads drawer query parameters such as /hub?drawer=reports and /hub?drawer=recordings."),
    ])

    doc.add_heading("16. Data Persistence", level=1)
    doc.add_paragraph(
        "This project is configured for a complete local demo, not a hosted multi-user database. Data is persisted in simple local files so the project can run without cloud setup."
    )
    add_matrix(
        doc,
        ["Data", "Storage Location", "Notes"],
        [
            ("Cases/sessions/reports metadata", "uploads/data/store.json", "Written through persistAll in server/src/db.js."),
            ("Recordings", "uploads/recordings", "Uploaded WebM evidence files."),
            ("PDF reports", "uploads/reports", "Generated report files."),
            ("Audit log", "uploads/data/store.json", "Most recent entries retained by local audit store."),
        ],
        [2500, 3100, 3760],
    )

    doc.add_heading("17. Setup and Run Instructions", level=1)
    add_code_block(doc, [
        "cd C:\\Users\\Admin\\Downloads\\altheria-fts-final-production-corrected-2026-05-27\\altheria-v3-fixed",
        "npm install",
        "cd server",
        "npm install",
        "cd ..\\client",
        "npm install",
        "cd ..",
        "npm run dev",
    ])
    doc.add_paragraph(
        "The root dev script starts the Express server and Vite client together. The backend defaults to http://localhost:4000/api and the Vite client normally opens at http://localhost:5173."
    )
    add_label_table(doc, [
        ("Build command", "cd client && npm run build"),
        ("Backend command", "cd server && npm run dev"),
        ("Frontend command", "cd client && npm run dev"),
        ("Local demo login", "Badge 8829 with password123, or ADMIN with admin123."),
    ])

    doc.add_heading("18. Browser Permission Checklist", level=1)
    add_bullets(doc, [
        "Use Chrome or Edge for best SpeechRecognition, SpeechSynthesis, MediaRecorder, microphone, and camera support.",
        "Allow microphone access when the browser asks.",
        "Allow camera access when the browser asks.",
        "Check Windows Settings > Privacy & security > Camera and Microphone if browser access is blocked.",
        "Use localhost or HTTPS. Browser media APIs are restricted on insecure origins.",
        "If camera preview fails, use Retry Camera in the assessment screen after confirming permissions.",
    ])

    doc.add_heading("19. Testing Checklist", level=1)
    add_numbered(doc, [
        "Start server and client with npm run dev.",
        "Login with a local demo investigator.",
        "Open or register a case.",
        "Start assessment in English and allow microphone/camera permissions.",
        "Confirm the first question speaks aloud.",
        "Answer by voice and wait five seconds; confirm answer submits automatically.",
        "Confirm the next question also speaks aloud and the system listens again.",
        "Repeat the same test in Hindi.",
        "End the assessment and confirm recordings upload.",
        "Open Reports, enter Case ID, and confirm only that case's reports appear.",
        "Open Recordings, enter Case ID, and confirm only that case's recordings appear.",
        "Download the PDF report and confirm voice analysis, camera analysis, answers, transcript, and recording evidence sections are present.",
    ])

    doc.add_heading("20. Important Limitations and Production Notes", level=1)
    add_bullets(doc, [
        "Camera access cannot be forced by code if Windows, the browser, or another application blocks the webcam.",
        "Browser SpeechRecognition support varies. Chrome/Edge are recommended; Firefox support is limited.",
        "Browser-generated text-to-speech audio is not always included in microphone recordings, so questions are preserved through transcript/report metadata.",
        "The current storage layer is local JSON and upload folders. For production, migrate to a real database and object storage.",
        "Set a strong JWT_SECRET in production and do not rely on local demo credentials.",
        "Use HTTPS in production so camera and microphone APIs work reliably.",
        "The behavioral and voice metrics are project indicators. They should be treated as decision-support signals, not definitive forensic proof.",
    ])

    doc.add_heading("21. Main Code Files Explained", level=1)
    add_matrix(
        doc,
        ["File", "Role", "Key Responsibilities"],
        [
            ("client/src/components/AssessmentEnclave.jsx", "Assessment UI", "TTS, STT, silence detection, camera, recordings, metadata, report generation, hide panels."),
            ("client/src/components/ReportsPanel.jsx", "History UI", "Case ID search, report-only view, recording-only view, evidence metadata display."),
            ("client/src/components/CommandHub.jsx", "Main workspace", "Drawer navigation and query-driven reports/recordings opening."),
            ("client/src/services/api.js", "Frontend API client", "Auth headers, REST calls, recording upload, report download URLs."),
            ("server/src/index.js", "Backend entry", "Middleware, routes, rate limits, health check, WebSocket server."),
            ("server/src/db.js", "Persistence", "Seed users/cases, maps, JSON persistence, audit append."),
            ("server/src/routes/interrogation.js", "Assessment API", "Opening questions, answers, transcript, session analysis flow."),
            ("server/src/routes/recordings.js", "Evidence API", "Multipart upload, list, stream/download recordings."),
            ("server/src/routes/reports.js", "Report API", "Generate/list/download reports."),
            ("server/src/services/aiInterrogator.js", "Question and text analysis", "Language profiles, local AI fallback, stress/sentiment/forensic scoring."),
            ("server/src/services/pdfReportService.js", "PDF report builder", "Case details, answer analysis, voice/camera sections, transcript, recording archive."),
        ],
        [3300, 1800, 4260],
    )

    doc.add_heading("22. Final Implemented Feature Summary", level=1)
    add_bullets(doc, [
        "Hands-free assessment flow works through spoken questions, spoken answers, and five-second silence submission.",
        "English and Hindi assessment modes are supported through browser language codes and backend language profiles.",
        "Voice recording evidence is uploaded and linked to the session and case.",
        "Camera capture and behavior metrics are attempted through browser camera APIs, with retry support.",
        "Reports include separate voice analysis, camera analysis, answer analysis, transcript, and evidence archive sections.",
        "Reports and recordings are case-filtered by Case ID and displayed in separate panels.",
        "Hide Panels in assessment hides camera/analysis panels while keeping the question-answer flow visible.",
    ])

    add_callout(
        doc,
        "Project status",
        "The codebase is implemented as a local full-stack project. The remaining real-world dependency is the browser/OS permission environment for microphone and camera hardware, which must be allowed on the user's machine.",
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
